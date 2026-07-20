import fitz
import os
from typing import List, Dict, Any

try:
    import docx
except ImportError:
    docx = None

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extrae todo el texto plano de un documento PDF de emergencia."""
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"El archivo {pdf_path} no existe.")
    
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        text += page.get_text() + "\n"
    doc.close()
    return text.strip()

def extract_text_from_docx(docx_path: str) -> str:
    """Extrae todo el texto de un documento de Word (.docx)."""
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"El archivo {docx_path} no existe.")
    if not docx:
        raise RuntimeError("La librería 'python-docx' no está instalada.")
    
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def extract_text_from_file(file_path: str) -> str:
    """Extrae texto automáticamente según la extensión del archivo (PDF, DOCX, TXT)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".txt"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext in [".jpg", ".jpeg", ".png", ".webp"]:
        return f"[Evidencia Fotográfica Adjunta: {os.path.basename(file_path)}]"
    else:
        return ""